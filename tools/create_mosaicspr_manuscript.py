from __future__ import annotations

import os
import re
import subprocess
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/manuscripts/MosaicSPR-methods.md"
BUILD = ROOT / "docs/manuscripts/build"
OUTPUT = BUILD / "MosaicSPR-methods-manuscript.docx"
REFERENCE = BUILD / "MosaicSPR-reference.docx"
TEMP_MD = BUILD / "MosaicSPR-pandoc.md"
PIPELINE = BUILD / "MosaicSPR-pipeline.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
INK = "202A30"
MUTED = "66747C"
GOLD = "7A5A00"
TABLE_FILL = "F4F6F9"
RULE = "D4DCE1"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None,
             bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_font(style, name: str, size: float, color: str = INK, bold: bool | None = None,
               italic: bool | None = None) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def style_paragraph(style, before: float, after: float, line: float,
                    alignment: WD_ALIGN_PARAGRAPH | None = None,
                    keep_with_next: bool = False) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if alignment is not None:
        fmt.alignment = alignment
    fmt.keep_with_next = keep_with_next
    fmt.widow_control = True


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    style_font(normal, "Calibri", 11, INK)
    style_paragraph(normal, 0, 8, 1.333, WD_ALIGN_PARAGRAPH.JUSTIFY)

    for name in ["Body Text", "First Paragraph"]:
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style_font(styles[name], "Calibri", 11, INK)
        style_paragraph(styles[name], 0, 8, 1.333, WD_ALIGN_PARAGRAPH.JUSTIFY)

    title = styles["Title"]
    style_font(title, "Calibri", 30, NAVY, bold=True)
    style_paragraph(title, 105, 8, 1.0, WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)
    subtitle = styles["Subtitle"]
    style_font(subtitle, "Calibri", 15, DARK_BLUE)
    style_paragraph(subtitle, 0, 28, 1.0, WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style_font(styles[name], "Calibri", size, color, bold=True)
        style_paragraph(styles[name], before, after, 1.0, WD_ALIGN_PARAGRAPH.LEFT, keep_with_next=True)

    for name in ["Author", "Date"]:
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style_font(styles["Author"], "Calibri", 10.5, MUTED, italic=True)
    style_paragraph(styles["Author"], 0, 5, 1.0, WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)
    style_font(styles["Date"], "Calibri", 12, NAVY, bold=True)
    style_paragraph(styles["Date"], 0, 4, 1.0, WD_ALIGN_PARAGRAPH.CENTER)

    for name in ["List Bullet", "List Number"]:
        style_font(styles[name], "Calibri", 11, INK)
        style_paragraph(styles[name], 0, 4, 1.208, WD_ALIGN_PARAGRAPH.LEFT)
        styles[name].paragraph_format.left_indent = Inches(0.375)
        styles[name].paragraph_format.first_line_indent = Inches(-0.194)

    style_font(styles["Caption"], "Calibri", 9.5, MUTED, italic=True)
    style_paragraph(styles["Caption"], 4, 8, 1.0, WD_ALIGN_PARAGRAPH.CENTER)

    if "Kicker" not in styles:
        styles.add_style("Kicker", WD_STYLE_TYPE.PARAGRAPH)
    style_font(styles["Kicker"], "Calibri", 10, GOLD, bold=True)
    style_paragraph(styles["Kicker"], 0, 18, 1.0, WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def paragraph_bottom_border(paragraph, color: str = RULE, size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_page_field(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    set_font(run, size=8.5, color=MUTED)
    for item in [begin, instr, separate, text, end]:
        run._r.append(item)


def configure_furniture(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = paragraph.add_run("MOSAICSPR")
    set_font(left, size=8.5, color=NAVY, bold=True)
    paragraph.add_run("\t")
    right = paragraph.add_run("METHODS MANUSCRIPT DRAFT")
    set_font(right, size=8.5, color=MUTED)
    paragraph_bottom_border(paragraph)

    footer = section.footer
    footer.is_linked_to_previous = False
    fpara = footer.paragraphs[0]
    fpara.clear()
    fpara.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fpara.paragraph_format.space_before = Pt(2)
    label = fpara.add_run("MosaicSPR  ·  ")
    set_font(label, size=8.5, color=MUTED)
    add_page_field(fpara)


def create_reference() -> None:
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        configure_furniture(section)

    doc.save(REFERENCE)


def create_pipeline_figure() -> None:
    labels = [
        ("Alignment encoding", "canonical bytes + bit planes"),
        ("Non-binding region proposals", "pair-covered informative triplets + overlap safety bank"),
        ("FastTree seed family", "global, segment, pair, triplet and overlapping windows"),
        ("Connected unrooted-SPR graph", "complete one-SPR neighborhoods + priced column generation"),
        ("Exact fixed-graph decoder", "minimum-duration runs + movable master + executable edit tape"),
    ]
    colors = ["#EAF3F7", "#E9F5F1", "#F7F1E5", "#EEEAF8", "#EAF1FB"]
    fig, ax = plt.subplots(figsize=(6.25, 6.15), dpi=220)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 11.5)
    ax.axis("off")
    centers = [10.25, 8.15, 6.05, 3.95, 1.85]
    for index, ((title, detail), y) in enumerate(zip(labels, centers)):
        box = FancyBboxPatch((0.65, y - 0.72), 8.7, 1.35,
                             boxstyle="round,pad=0.025,rounding_size=0.12",
                             linewidth=1.2, edgecolor="#54727F", facecolor=colors[index])
        ax.add_patch(box)
        ax.text(1.02, y + 0.17, f"{index + 1}", ha="center", va="center", fontsize=10,
                fontweight="bold", color="#FFFFFF",
                bbox=dict(boxstyle="circle,pad=0.25", facecolor="#2E6F75", edgecolor="none"))
        ax.text(1.55, y + 0.18, title, ha="left", va="center", fontsize=11.2,
                fontweight="bold", color="#203748")
        ax.text(1.55, y - 0.25, detail, ha="left", va="center", fontsize=8.8, color="#53636B")
        if index + 1 < len(labels):
            ax.annotate("", xy=(5, centers[index + 1] + 0.75), xytext=(5, y - 0.75),
                        arrowprops=dict(arrowstyle="-|>", color="#6A7E86", lw=1.3))
    ax.text(5, 11.15, "MosaicSPR separates proposal generation from final reconstruction",
            ha="center", va="center", fontsize=12.2, fontweight="bold", color="#203748")
    fig.savefig(PIPELINE, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def prepare_markdown() -> None:
    source = SOURCE.read_text(encoding="utf-8")
    start = source.index("## Abstract")
    body = source[start:]

    def lower_heading(match: re.Match[str]) -> str:
        hashes = match.group(1)
        return "#" * max(1, len(hashes) - 1) + " "

    body = re.sub(r"^(#{2,3})\s+", lower_heading, body, flags=re.MULTILINE)
    figure = (
        f"![]({PIPELINE.as_posix()}){{ width=78% }}\n\n"
        "*Figure 1. MosaicSPR computational workflow. Breakpoint proposals and FastTree fits seed the search but do not constrain the exact fixed-graph segmentation.*\n\n"
    )
    body = body.replace("The key separation is between", figure + "The key separation is between", 1)
    metadata = """---
title: "MosaicSPR"
subtitle: "Fast reconstruction of mosaic phylogenies as explicit subtree-prune-and-regraft histories"
author: "Methods manuscript draft · Author list and affiliations to be supplied"
date: "August 2026"
---

"""
    TEMP_MD.write_text(metadata + body, encoding="utf-8")


def patch_numbering(doc: Document) -> None:
    numbering = doc.part.numbering_part.element
    for level in numbering.findall(".//" + qn("w:lvl")):
        p_pr = level.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            level.append(p_pr)
        tabs = p_pr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            p_pr.append(tabs)
        for old in list(tabs):
            tabs.remove(old)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        indent = p_pr.find(qn("w:ind"))
        if indent is None:
            indent = OxmlElement("w:ind")
            p_pr.append(indent)
        indent.set(qn("w:left"), "540")
        indent.set(qn("w:hanging"), "279")


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths: list[int]) -> None:
    assert sum(widths) == 9120
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    # Pandoc writes references to undefined "Table" and "Compact" styles.
    # LibreOffice interprets those references as a floating empty grid followed
    # by the cell text in normal flow, so remove the table style and normalize
    # every cell paragraph to the document's real Normal style.
    tbl_style = tbl_pr.find(qn("w:tblStyle"))
    if tbl_style is not None:
        tbl_pr.remove(tbl_style)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9120")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "C8D2D9")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for col_index, cell in enumerate(row.cells):
            width = widths[min(col_index, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            if row_index == 0:
                shade_cell(cell, TABLE_FILL)
            for paragraph in cell.paragraphs:
                paragraph.style = "Normal"
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.08
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_font(run, size=9.3, color=INK, bold=True if row_index == 0 else None)


def postprocess() -> None:
    doc = Document(OUTPUT)
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        configure_furniture(section)

    for shape in doc.inline_shapes:
        doc_pr = shape._inline.docPr
        doc_pr.set("descr", "MosaicSPR computational workflow from alignment encoding through exact fixed-graph decoding")
        doc_pr.set("title", "MosaicSPR computational workflow")

    title = next((p for p in doc.paragraphs if p.style.name == "Title"), None)
    if title is None:
        raise RuntimeError("Pandoc did not create a Title paragraph.")
    kicker = title.insert_paragraph_before("METHODS MANUSCRIPT")
    kicker.style = doc.styles["Kicker"]
    for run in kicker.runs:
        set_font(run, size=10, color=GOLD, bold=True)

    date = next((p for p in doc.paragraphs if p.text.strip() == "August 2026"), None)
    if date is None:
        raise RuntimeError("Pandoc did not create the expected cover date.")
    date.add_run().add_break(WD_BREAK.PAGE)

    # Keep headings with their following paragraph and normalize list spacing.
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        p_pr = paragraph._p.pPr
        if p_pr is not None and p_pr.numPr is not None:
            paragraph.paragraph_format.left_indent = Inches(0.375)
            paragraph.paragraph_format.first_line_indent = Inches(-0.194)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.208
        if paragraph.text.startswith("Figure 1."):
            paragraph.style = doc.styles["Caption"]

    # The methods manuscript contains one three-column parameter table.
    for table in doc.tables:
        if len(table.columns) == 3:
            set_table_geometry(table, [2220, 1840, 5060])
        elif len(table.columns) == 2:
            set_table_geometry(table, [2740, 6380])
        else:
            count = max(1, len(table.columns))
            base = 9120 // count
            widths = [base] * count
            widths[-1] += 9120 - sum(widths)
            set_table_geometry(table, widths)

    patch_numbering(doc)
    doc.core_properties.title = "MosaicSPR methods manuscript"
    doc.core_properties.subject = "Technical methods and practical implementation details"
    doc.core_properties.author = "Author list to be supplied"
    doc.core_properties.keywords = "MosaicSPR, recombination, phylogenetics, subtree prune and regraft, EvoOnline"
    doc.save(OUTPUT)


def audit() -> None:
    doc = Document(OUTPUT)
    for section in doc.sections:
        assert round(section.page_width.inches, 3) == 8.5
        assert round(section.page_height.inches, 3) == 11.0
        assert all(round(value.inches, 3) == 1.0 for value in [section.top_margin, section.right_margin, section.bottom_margin, section.left_margin])
        assert round(section.header_distance.inches, 3) == 0.492
        assert round(section.footer_distance.inches, 3) == 0.492
    normal = doc.styles["Normal"]
    assert normal.font.name == "Calibri" and round(normal.font.size.pt, 1) == 11.0
    assert round(float(normal.paragraph_format.line_spacing), 3) == 1.333
    assert round(normal.paragraph_format.space_after.pt, 1) == 8.0
    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        assert tbl_w is not None and tbl_w.get(qn("w:w")) == "9120" and tbl_w.get(qn("w:type")) == "dxa"
        assert tbl_ind is not None and tbl_ind.get(qn("w:w")) == "120"
        grid = [int(item.get(qn("w:w"))) for item in table._tbl.tblGrid]
        assert sum(grid) == 9120
        for row in table.rows:
            widths = [int(cell._tc.get_or_add_tcPr().find(qn("w:tcW")).get(qn("w:w"))) for cell in row.cells]
            assert widths == grid
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "##" not in text
    assert "Methods manuscript draft" in text
    assert len(doc.paragraphs) > 100
    print(f"AUDIT OK: {OUTPUT} · {len(doc.paragraphs)} paragraphs · {len(doc.tables)} tables")


def main() -> None:
    BUILD.mkdir(parents=True, exist_ok=True)
    create_pipeline_figure()
    create_reference()
    prepare_markdown()
    subprocess.run([
        "pandoc", str(TEMP_MD),
        "--from=markdown+pipe_tables+tex_math_dollars+tex_math_single_backslash+smart",
        "--to=docx",
        f"--reference-doc={REFERENCE}",
        f"--output={OUTPUT}",
    ], check=True)
    postprocess()
    audit()


if __name__ == "__main__":
    main()
